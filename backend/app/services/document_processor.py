import asyncio
import os
import uuid
from typing import List, Dict, Any, Tuple
from pathlib import Path
import mimetypes
from PIL import Image
import io

# Text processing
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.services.llm_service import gemini_service
from app.models.document import Document, DocumentChunk
from sqlalchemy.orm import Session

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_document(
        self, 
        file_path: str, 
        filename: str, 
        db: Session,
        user_id: Optional[int] = None
    ) -> Document:
        """Process a document and create chunks with embeddings"""
        
        # Determine file type and content
        file_type = self._get_file_type(filename)
        file_size = os.path.getsize(file_path)
        
        # Create document record
        document = Document(
            filename=str(uuid.uuid4()) + Path(filename).suffix,
            original_name=filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            content_type=self._determine_content_type(file_type),
            user_id=user_id
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        try:
            # Extract content based on file type
            if file_type in ['.pdf', '.txt', '.md', '.docx']:
                await self._process_text_document(document, db)
            elif file_type in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                await self._process_image_document(document, db)
            
            # Mark as processed
            document.processed = True
            db.commit()
            
        except Exception as e:
            db.rollback()
            raise Exception(f"Error processing document {filename}: {str(e)}")
        
        return document
    
    async def _process_text_document(self, document: Document, db: Session):
        """Process text-based documents"""
        
        # Extract text content
        text_content = self._extract_text(document.file_path, document.file_type)
        
        if not text_content.strip():
            raise Exception("No text content found in document")
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text_content)
        
        # Process chunks in batches to avoid API rate limits
        batch_size = 5
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            await self._process_chunk_batch(document, batch, i, db)
    
    async def _process_image_document(self, document: Document, db: Session):
        """Process image documents"""
        
        # Read image
        with open(document.file_path, 'rb') as f:
            image_data = f.read()
        
        # Analyze image with Gemini Vision
        analysis = await gemini_service.analyze_image(
            image_data, 
            "Describe this image in detail, including any text, objects, people, scenes, and relevant context."
        )
        
        # Generate embedding for the analysis
        embedding = await gemini_service.generate_embeddings(analysis)
        
        # Create single chunk for image
        chunk = DocumentChunk(
            document_id=document.id,
            chunk_index=0,
            content=analysis,
            metadata=json.dumps({
                "type": "image_analysis",
                "original_filename": document.original_name,
                "file_type": document.file_type
            })
        )
        
        db.add(chunk)
        db.commit()
        db.refresh(chunk)
        
        # Add to vector store
        from app.services.vector_store import vector_store
        embedding_id = vector_store.add_document(
            content=analysis,
            embedding=embedding,
            metadata={
                "document_id": document.id,
                "chunk_index": 0,
                "type": "image_analysis",
                "original_filename": document.original_name
            },
            document_id=document.id,
            chunk_index=0
        )
        
        chunk.embedding_id = embedding_id
        db.commit()
    
    async def _process_chunk_batch(
        self, 
        document: Document, 
        chunk_texts: List[str], 
        start_index: int, 
        db: Session
    ):
        """Process a batch of text chunks"""
        
        # Generate embeddings for all chunks in batch
        embedding_tasks = [
            gemini_service.generate_embeddings(chunk_text) 
            for chunk_text in chunk_texts
        ]
        embeddings = await asyncio.gather(*embedding_tasks)
        
        # Create chunk records and add to vector store
        for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
            chunk_index = start_index + i
            
            # Create database record
            chunk = DocumentChunk(
                document_id=document.id,
                chunk_index=chunk_index,
                content=chunk_text,
                metadata=json.dumps({
                    "type": "text_chunk",
                    "original_filename": document.original_name,
                    "file_type": document.file_type,
                    "chunk_size": len(chunk_text)
                })
            )
            
            db.add(chunk)
            db.commit()
            db.refresh(chunk)
            
            # Add to vector store
            from app.services.vector_store import vector_store
            embedding_id = vector_store.add_document(
                content=chunk_text,
                embedding=embedding,
                metadata={
                    "document_id": document.id,
                    "chunk_index": chunk_index,
                    "type": "text_chunk",
                    "original_filename": document.original_name
                },
                document_id=document.id,
                chunk_index=chunk_index
            )
            
            chunk.embedding_id = embedding_id
            db.commit()
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from different file types"""
        
        if file_type == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_type == '.docx':
            return self._extract_docx_text(file_path)
        elif file_type in ['.txt', '.md']:
            return self._extract_plain_text(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_plain_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _get_file_type(self, filename: str) -> str:
        """Get file extension"""
        return Path(filename).suffix.lower()
    
    def _determine_content_type(self, file_type: str) -> str:
        """Determine content type based on file extension"""
        if file_type in ['.pdf', '.txt', '.md', '.docx']:
            return 'text'
        elif file_type in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
            return 'image'
        else:
            return 'unknown'